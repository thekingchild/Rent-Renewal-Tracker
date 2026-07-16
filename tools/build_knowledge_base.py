from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Rent Renewal Tracker User Guide.md"
OUTPUT = ROOT / "docs" / "Rent Renewal Tracker Knowledge Base and User Manual.docx"

NAVY = RGBColor(20, 47, 66)
BLUE = RGBColor(36, 99, 130)
TEAL = RGBColor(20, 122, 128)
GOLD = RGBColor(190, 132, 42)
INK = RGBColor(38, 48, 56)
MUTED = RGBColor(95, 108, 118)
PALE_BLUE = "E8EEF5"
PALE_TEAL = "E7F3F2"
PALE_GOLD = "FBF3E2"
GRID = "CBD5DC"
WHITE = RGBColor(255, 255, 255)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    if table.rows:
        set_repeat_table_header(table.rows[0])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=None, color=INK, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_rich_text(paragraph, text, size=10.5, color=INK):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_run(paragraph.add_run(text[pos:match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("`"):
            set_run(paragraph.add_run(token[1:-1]), size=size - 0.3, color=BLUE, font="Consolas")
        else:
            set_run(paragraph.add_run(token[2:-2]), size=size, color=color, bold=True)
        pos = match.end()
    if pos < len(text):
        set_run(paragraph.add_run(text[pos:]), size=size, color=color)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.38)
    sec.footer_distance = Inches(0.38)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 13, MUTED, 0, 10),
        ("Heading 1", 17, NAVY, 18, 8),
        ("Heading 2", 13.5, BLUE, 14, 6),
        ("Heading 3", 11.5, TEAL, 10, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.add_run("RENT RENEWAL TRACKER  /  KNOWLEDGE BASE"), size=8.5, color=MUTED, bold=True)
    footer = sec.footer
    add_page_number(footer.paragraphs[0])


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(kicker.add_run("APPLICATION KNOWLEDGE BASE"), size=10.5, color=GOLD, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(title.add_run("Rent Renewal Tracker"), size=30, color=NAVY, bold=True)
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(sub.add_run("Complete User Manual, Operations Guide, Permissions Reference and Field Catalogue"), size=14, color=BLUE)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Frappe Framework v16 application  |  Application version 0.2.0"), size=10, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Prepared 16 July 2026"), size=10, color=MUTED)
    for _ in range(5):
        doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    set_cell_fill(callout.cell(0, 0), PALE_TEAL)
    p = callout.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_text(p, "Purpose: enable a new user, approver, auditor, or administrator to understand and operate every user-facing feature of the application without relying on undocumented institutional knowledge.", size=10.5)
    doc.add_page_break()


def add_front_matter(doc):
    doc.add_heading("How to use this knowledge base", level=1)
    p = doc.add_paragraph()
    add_rich_text(p, "Start with Sections 1-6 for orientation and setup. Operational users should then study leases, documents, payments, and lifecycle requests. Approvers should focus on the workflow and permissions chapters. Administrators and auditors should also study reminders, automation, reporting, security controls, and the field catalogue.")
    add_callout(doc, "Authoritative scope", "This manual is derived from the application source in this workspace as of 16 July 2026. It documents implemented behavior. Site-specific Frappe configuration, custom roles, email accounts, print formats, and future code changes may alter what an individual user sees.", PALE_GOLD)
    doc.add_heading("Contents at a glance", level=2)
    contents = [
        "Purpose, scope, roles, access, and workspace navigation",
        "Implementation setup and readiness checks",
        "Master data: departments, properties, and landlords",
        "Lease creation, submission, status logic, calendar, and lifecycle actions",
        "Private lease documents, expiry monitoring, and immutable revisions",
        "Rent schedules, payment evidence, status derivation, and calendar",
        "Renewal and termination workflow, approvals, execution, and successor leases",
        "Reminder policies, recipient resolution, delivery, retries, logs, and weekly digest",
        "Dashboard cards, My Actions, reports, charts, filters, and interpretation",
        "Security model, role permissions, confidentiality clearance, and department scoping",
        "Troubleshooting, operating controls, glossary, and complete field catalogue",
    ]
    for item in contents:
        p = doc.add_paragraph(style="List Bullet")
        add_rich_text(p, item)
    doc.add_page_break()


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_fill(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    set_run(p.add_run(f"{label}: "), size=10.3, color=NAVY, bold=True)
    add_rich_text(p, text, size=10.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def markdown_to_doc(doc, markdown):
    lines = markdown.splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("Version:") or stripped.startswith("Prepared:"):
            continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\. ", "", stripped))
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            add_rich_text(p, stripped)


SUPPLEMENT = r"""
## 21. Dashboard, Personal Work Queue, and Visual Status Cues

### 21.1 Dashboard cards and drill-down behavior

The workspace contains permission-aware cards for Expiring in 30 Days, Expiring in 60 Days, Expiring in 90 Days, Renewals Waiting for Me, Overdue Rent Obligations, Failed Reminders, and Annual Rent Exposure. Selecting a card opens the corresponding report with relevant filters. Annual Rent Exposure uses the configured default currency. If no default exists and the visible data contains more than one currency, it displays Multiple currencies instead of adding incomparable amounts.

### 21.2 Dashboard charts

Lease Expiry Outlook uses the Upcoming Expiries report chart. Renewal Workflow uses the Renewal Pipeline report chart. Both inherit the current user's record-level access; dashboard counts are not intended to reveal inaccessible leases.

### 21.3 My Actions

My Actions consolidates Renewal Approval, Overdue Payment, Lease Expiry, Document Expiry, and Failed Reminder work. Filter by Action Type or Priority. Critical items sort before High, Medium, and Low; within a priority, the earliest event date sorts first. Renewal approvals become High when older than seven days. Overdue payments become Critical when more than 30 days late. Lease expiries are High within 30 days, Medium within 60, and Low through 90 days. Expired documents and failed reminders are Critical; documents expiring soon are High.

### 21.4 List and calendar indicators

Lease, Rent Schedule, Renewal Request, and Lease Document list views use colored indicators. Lease and Rent Schedule also provide calendar views. Lease calendar events span Start Date through End Date; Rent Schedule calendar events appear on Due Date. A Restricted lease displays a red access warning. Lease forms show the next contractual action. Renewal forms highlight changed proposed terms and the role currently expected to act. Lease Document forms warn when a document is Expiring Soon or Expired.

## 22. Security, Permissions, and Information Boundaries

### 22.1 Layered authorization model

Access is the intersection of four controls: application role, DocType permission, record assignment or department scope, and confidentiality clearance. A user may possess read permission for the Lease DocType and still see no record because the lease is outside the user's assignment, department permission, or clearance level. Guest users receive no application access.

### 22.2 Unrestricted operational roles

Administrator, System Manager, Rent Renewal System Manager, and Lease Administrator bypass the lease assignment and confidentiality filters. This is operationally powerful; assign these roles only to trusted administrators. It does not change every DocType's create/write permission: the DocType matrix still determines what each role can do.

For Lease Document specifically, the default create grant belongs to Rent Renewal System Manager, Lease Administrator, and Responsible Officer. Administrator retains the Frappe permission bypass. System Manager alone bypasses the record-level filter but cannot create a Lease Document because the default Lease Document permission matrix does not grant that role create permission.

### 22.3 Assigned-user and department access

For scoped users, a lease is visible when the user is the Responsible Officer, Contract Owner, or Backup Officer, or when a Frappe User Permission grants that user the lease's Lease Department. The User Permission must allow Lease Department, match Responsible Department, and either be globally applicable or applicable to Lease. The same rule flows down to Renewal Requests, Rent Schedules, Lease Documents, and Reminder Logs through their Lease link. Lease Documents add their own confidentiality check after the parent Lease check.

When a new Lease is saved, the same scope is evaluated from its proposed values before insertion. Unrestricted operational roles may create without assignment. A Responsible Officer must assign themselves as Responsible Officer, Contract Owner, or Backup Officer, or hold the matching Lease Department User Permission, and the selected classification must be within their clearance. Related records must link to a Lease the user is authorized to access.

### 22.4 Confidentiality clearance

Lease Viewer can see Public and Internal leases. Responsible Officer, Department Head, and Finance Approver can see Public, Internal, and Confidential leases. Legal Approver, Management Approver, and Lease Auditor can see Public, Internal, Confidential, and Restricted leases. A user with multiple roles receives the union of those clearances. Lease Document access must pass this parent clearance and the document's Internal, Confidential, or Restricted classification; a user may see the Lease while a stricter linked document and private file remain hidden.

Role permissions are also combined. A user with Responsible Officer plus Legal Approver, Management Approver, or Lease Auditor receives Lease Document create permission from Responsible Officer and Restricted clearance from the additional role, but the user must still be assigned to the Lease or hold the matching Lease Department User Permission.

### 22.5 Workflow authorization and segregation of duties

An approver must have the role permitted for the current workflow transition and must be authorized for the lease. The requester cannot approve their own Department, Finance, Legal, or Management review stage. Return and Reject require a comment. Generic Frappe workflow email is disabled because it could attach a full PDF to every role holder; the application instead uses scoped notifications and reminder recipients.

For Renewal Request creation, the default create grant belongs to Rent Renewal System Manager, Lease Administrator, and Responsible Officer, while Administrator retains the Frappe permission bypass. System Manager alone and the review-only roles cannot create a Renewal Request. A Responsible Officer must pass the same assignment or Lease Department scope and confidentiality checks as the linked Lease.

Workflow authority is separate from DocType authority. Rent Renewal System Manager can create a Renewal Request but has no transition assigned to that role alone. Lease Administrator can submit a Draft and mark an Approved request executed, while Department Head, Finance Approver, Legal Approver, and Management Approver act only at their configured stages. Every scoped approver must also pass the linked-Lease authorization check.

### 22.6 Private files and audit integrity

Lease Documents accept only private files. The user must be able to read the File and either own an unattached upload, own a file attached to the lease/document, use a temporary upload, or have read access to the original attached record. Once attached, Frappe delegates File reads and downloads to the Lease Document permission decision. Existing document files cannot be replaced; a new revision requires create permission on Lease Document and write permission on the current revision. Reminder Logs cannot be edited manually and cannot be deleted except during uninstall.

## 23. Lifecycle Actions: Renewal and Termination

### 23.1 Start from a lease

On a submitted lease in Active, Expiring Soon, or Expired status, open Lifecycle and select Start Renewal or Start Termination. The application creates a Renewal Request with recommendation Renew or Terminate and opens it. If an open request of the same type already exists, the existing request opens. If the open request is of the other type, complete or cancel it first.

The Lease lifecycle action is the recommended entry point because it applies the submitted/status eligibility checks. Direct Add Renewal Request creation remains available to create-enabled roles but does not currently apply that same eligibility gate. The Lease form may also display the lifecycle button to a read-only user; server-side create permission still prevents unauthorized creation.

### 23.2 Termination-specific behavior

Termination uses the same approval chain. Outside Draft, Termination Effective Date and Termination Reason are mandatory. The parent lease displays Termination in Progress during the cycle. Before Mark Executed, attach a current private Approval document dated and linked to the lifecycle request. Completion marks the lease Terminated and does not create a successor lease.

### 23.3 Renewal-specific execution

For Renew, Renegotiate, or Relocate, the approved request requires proposed property, start date, end date, currency, rent basis, and payment frequency before leaving Draft. Before execution, attach a current private Renewal Letter with Document Date and link it to the same Lease and exact Renewal Request. Only Lease Administrator can perform the configured Mark Executed transition. Completion creates one successor lease, links predecessor and successor, copies ownership, contacts, landlord, classification, and selected financial structure, and applies approved proposed terms. Application validation permits one non-cancelled successor while retaining cancelled and amended revision history.

### 23.4 Cancellation and rejection

Reject is available only at Management Approval and stores the workflow comment as Rejection Reason. Canceling an open request clears the open-cycle key and returns the lease to Active or Expired based on its End Date. Rejection similarly restores the lease's operational status while retaining the request and decision history as evidence.

## 24. Lease Document Lifecycle and Revision Control

### 24.1 Supported files and size limits

Allowed extensions are CSV, DOCX, JPEG, JPG, PDF, PNG, TXT, and XLSX. The file must be private and must remain within the site's configured maximum file size.

### 24.2 Status calculation

No Expiry Date is used when Expiry Date is blank. Current means the document is beyond the configured document-expiry warning window. Expiring Soon means the remaining days are at or below Document Expiring Soon Threshold. Expired means the date has passed. Superseded means a later revision exists. Daily automation refreshes these values.

### 24.3 Creating a revision

Open the current Lease Document and select Create Revision. The action is shown only when the user can create Lease Documents and write the current revision. The new record inherits lease, lifecycle request, title, category, dates, and confidentiality. Upload a different private file and enter Revision Reason. Frappe assigns the new record name before Document Family ID and Revision number are calculated. Saving uses normal permission, validation, and change-tracking hooks to mark the previous revision Superseded and clear its expiry-attention value. Revisions must remain on the same lease and must branch only from the current revision; direct API calls receive the same checks.

## 25. Payment Evidence and Schedule Integrity

Partially Paid and Paid schedules require both Payment Reference and Paid On. Waived schedules require a waiver reason in Notes. Schedule periods cannot overlap another non-cancelled schedule for the same lease, and must fall within the lease term. Total Due equals Base Rent plus Service Charge plus Tax. The currency must match the lease. After a schedule is submitted, daily automation may refresh Schedule Status without changing submitted commercial values.

## 26. Reminder Engine, Delivery Audit, and Digest

### 26.1 Eligibility and threshold selection

The daily engine evaluates leases with an End Date, excluding Draft, Renewed, and Terminated lease statuses and Completed or Not Renewing renewal statuses. It uses the lease-specific Reminder Policy when set, otherwise the default policy. It selects a due enabled threshold that has not already been logged for the lease, renewal cycle, policy, channel, and recipient. This deduplication prevents repeat sends for the same event.

### 26.2 Recipient resolution

Lease User Field supports responsible_officer, backup_officer, and contract_owner. Lease Contact Type resolves only contacts of the selected type whose Receives Reminders box is selected. Role resolves all enabled users holding that role. Explicit User resolves one enabled user. Email uses a direct validated address. Scope limits a rule to All, Pre-Expiry, Expiry, or Overdue.

### 26.3 Delivery states and retries

Each delivery is reserved as Queued. A worker claims it as Sending so duplicate workers exit without sending a second message. Success becomes Sent, records Sent At and Message ID, and updates the lease's Last Reminder Date and Next Reminder Date. Failure becomes Failed, records the error and retry count, writes an application error log, notifies configured administrator recipients, and queues another attempt while the retry limit allows.

### 26.4 Auto-created renewal requests

When Create Renewal Request at First Threshold is enabled, the engine may create one request at the largest enabled pre-expiry threshold if the lease has no last request. It recommends Renegotiate, begins the proposed term the day after the current End Date, and uses approximately the same tenure length.

### 26.5 Weekly management digest

When enabled, the weekly job emails configured recipients a summary of leases expiring within 90 days, approvals older than seven days, overdue payments, failed reminders, and annual rent exposure grouped by currency. It includes links to My Actions and Renewal Pipeline. If enabled without recipients, the system logs an error instead of sending.

## 27. Administrative Readiness and Health Controls

Setup Readiness checks for at least one Lease Department, assigned users with core roles, a valid default Reminder Policy, an outgoing Email Account, at least one Property, and at least one Lease. Complete these items before depending on automated reminders. A technical installation health check also verifies the installed app, required DocTypes, roles, workflow states and actions, active workflow, reminder policy, workspace, reports, cards, and charts.

## 28. Core Role Permission Matrix
"""


def user_doctypes():
    base = ROOT / "rent_renewal_tracker" / "rent_renewal_tracker" / "doctype"
    docs = []
    for path in sorted(base.glob("*/*.json")):
        data = json.loads(path.read_text(encoding="utf-8"))
        if not data.get("istable"):
            docs.append(data)
    return docs


def add_permissions_matrix(doc, doctypes):
    roles = [
        "Rent Renewal System Manager", "Lease Administrator", "Responsible Officer",
        "Department Head", "Finance Approver", "Legal Approver", "Management Approver",
        "Lease Auditor", "Lease Viewer",
    ]
    key_docs = ["Lease", "Renewal Request", "Rent Schedule", "Lease Document", "Reminder Policy", "Reminder Log", "Rent Renewal Settings"]
    table = doc.add_table(rows=1, cols=1 + len(key_docs))
    widths = [1700] + [1094] * 6 + [1096]
    set_table_geometry(table, widths)
    headers = ["Role"] + key_docs
    for i, label in enumerate(headers):
        set_cell_fill(table.cell(0, i), "173B4D")
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(label), size=8, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    lookup = {d["name"]: d for d in doctypes}
    for role in roles:
        cells = table.add_row().cells
        set_run(cells[0].paragraphs[0].add_run(role), size=8.2, color=NAVY, bold=True)
        for idx, dt in enumerate(key_docs, 1):
            perms = [p for p in lookup[dt].get("permissions", []) if p.get("role") == role]
            flags = set()
            for perm in perms:
                for key, code in (("read", "R"), ("write", "W"), ("create", "C"), ("submit", "S"), ("cancel", "X"), ("delete", "D"), ("export", "E")):
                    if perm.get(key): flags.add(code)
            value = "/".join([c for c in ("R", "W", "C", "S", "X", "D", "E") if c in flags]) or "-"
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(value), size=8.2, color=INK)
        if len(table.rows) % 2 == 1:
            for c in cells: set_cell_fill(c, "F5F7F8")
    p = doc.add_paragraph()
    add_rich_text(p, "Legend: R read, W write, C create, S submit, X cancel, D delete, E export. Print/email permissions and master-register detail are omitted from this compact matrix but remain governed by the DocType configuration. Record-level scope and confidentiality still apply.", size=9.2, color=MUTED)

    doc.add_heading("28.1 Detailed DocType permissions", level=2)
    p = doc.add_paragraph()
    add_rich_text(p, "The following inventory expands the compact matrix and includes every configured role permission for every non-child application DocType. Duplicate permission rows for the same role are combined.")
    action_map = (
        ("read", "Read"), ("write", "Write"), ("create", "Create"),
        ("delete", "Delete"), ("submit", "Submit"), ("cancel", "Cancel"),
        ("amend", "Amend"), ("print", "Print"), ("email", "Email"),
        ("export", "Export"),
    )
    for data in sorted(doctypes, key=lambda d: d["name"]):
        doc.add_heading(data["name"], level=3)
        merged = {}
        for perm in data.get("permissions", []):
            role = perm.get("role")
            if not role:
                continue
            merged.setdefault(role, set())
            for key, label in action_map:
                if perm.get(key):
                    merged[role].add(label)
        if not merged:
            p = doc.add_paragraph()
            add_rich_text(p, "No direct application-role permissions are defined; access is inherited through the parent record or framework controls.", size=9.2, color=MUTED)
            continue
        table = doc.add_table(rows=1, cols=2)
        set_table_geometry(table, [3000, 6360])
        for i, label in enumerate(("Role", "Allowed actions")):
            set_cell_fill(table.cell(0, i), "245B73")
            set_run(table.cell(0, i).paragraphs[0].add_run(label), size=8.5, color=WHITE, bold=True)
        set_repeat_table_header(table.rows[0])
        for n, role in enumerate(roles):
            if role not in merged:
                continue
            cells = table.add_row().cells
            set_run(cells[0].paragraphs[0].add_run(role), size=8.4, color=NAVY, bold=True)
            set_run(cells[1].paragraphs[0].add_run(", ".join(label for _, label in action_map if label in merged[role])), size=8.4, color=INK)
            if n % 2:
                for c in cells: set_cell_fill(c, "F5F7F8")


def add_field_catalog(doc, doctypes):
    doc.add_page_break()
    doc.add_heading("29. Complete User-Facing Field Catalogue", level=1)
    p = doc.add_paragraph()
    add_rich_text(p, "This catalogue lists every non-layout field defined in the application's user-facing DocTypes and child tables. Required means the form definition requires a value. Read-only means the application calculates or controls the value. Conditional visibility is shown where implemented.")
    child_names = {"Lease Contact", "Reminder Threshold", "Reminder Recipient", "Renewal Decision"}
    ordered = sorted(doctypes, key=lambda d: (d["name"] not in ("Lease", "Renewal Request", "Rent Schedule", "Lease Document"), d["name"]))
    base = ROOT / "rent_renewal_tracker" / "rent_renewal_tracker" / "doctype"
    for path in sorted(base.glob("*/*.json")):
        data = json.loads(path.read_text(encoding="utf-8"))
        if data.get("name") in child_names:
            ordered.append(data)
    for data in ordered:
        doc.add_heading(data["name"], level=2)
        meta = []
        if data.get("autoname"): meta.append(f"Naming: {data['autoname']}")
        if data.get("issingle"): meta.append("Single settings record")
        if data.get("is_submittable"): meta.append("Submittable record")
        if data.get("istable"): meta.append("Child table")
        if meta:
            p = doc.add_paragraph()
            set_run(p.add_run(" | ".join(meta)), size=9.2, color=MUTED, italic=True)
        fields = [f for f in data.get("fields", []) if f.get("fieldtype") not in {"Section Break", "Column Break", "Tab Break", "HTML", "Heading"}]
        table = doc.add_table(rows=1, cols=4)
        set_table_geometry(table, [2200, 1350, 1300, 4510])
        for i, label in enumerate(("Field", "Type", "Control", "Options / behavior")):
            set_cell_fill(table.cell(0, i), "245B73")
            set_run(table.cell(0, i).paragraphs[0].add_run(label), size=8.5, color=WHITE, bold=True)
        set_repeat_table_header(table.rows[0])
        for n, f in enumerate(fields):
            cells = table.add_row().cells
            control = []
            if f.get("reqd"): control.append("Required")
            if f.get("read_only"): control.append("Read-only")
            if f.get("unique"): control.append("Unique")
            if f.get("depends_on"): control.append("Conditional")
            options = str(f.get("options") or "").replace("\n", "; ")
            behavior = []
            if options: behavior.append(options)
            if f.get("default") not in (None, ""): behavior.append(f"Default: {f['default']}")
            if f.get("depends_on"): behavior.append(f"Shown when: {f['depends_on']}")
            values = [f.get("label") or f.get("fieldname"), f.get("fieldtype") or "", ", ".join(control) or "Standard", ". ".join(behavior) or "-"]
            for i, value in enumerate(values):
                set_run(cells[i].paragraphs[0].add_run(value), size=8.2, color=INK, bold=(i == 0))
            if n % 2:
                for c in cells: set_cell_fill(c, "F5F7F8")


def add_final_sections(doc):
    doc.add_page_break()
    doc.add_heading("30. Glossary and Status Reference", level=1)
    entries = [
        ("DocType", "A Frappe record type, comparable to a business object or database-backed form."),
        ("Draft / Submitted / Cancelled", "Frappe document states 0, 1, and 2. Lease and Rent Schedule are submittable; workflow may submit Renewal Requests at Approved/Completed."),
        ("Lease status", "Operational condition of the lease: Draft, Active, Expiring Soon, Renewal in Progress, Termination in Progress, Renewed, Expired, or Terminated."),
        ("Renewal status", "Summary copied to the lease: Not Started, Draft, Pending Approval, Approved, Rejected, Completed, or Not Renewing."),
        ("Workflow state", "The detailed approval position on a Renewal Request."),
        ("Open cycle", "A non-final Renewal Request that reserves the lease so another lifecycle request cannot start."),
        ("Reminder threshold", "A configured number of days before End Date when an event becomes eligible."),
        ("Overdue cadence", "The minimum interval, in days, between overdue reminder events."),
        ("Confidentiality classification", "The lease-level access label Public, Internal, Confidential, or Restricted."),
        ("Document family", "All immutable revisions of the same Lease Document."),
        ("Successor lease", "The new lease created when a non-termination lifecycle request is completed."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2300, 7060])
    for i, label in enumerate(("Term", "Meaning")):
        set_cell_fill(table.cell(0, i), "173B4D")
        set_run(table.cell(0, i).paragraphs[0].add_run(label), size=9, color=WHITE, bold=True)
    for n, (term, meaning) in enumerate(entries):
        cells = table.add_row().cells
        set_run(cells[0].paragraphs[0].add_run(term), size=9, color=NAVY, bold=True)
        set_run(cells[1].paragraphs[0].add_run(meaning), size=9, color=INK)
        if n % 2: 
            for c in cells: set_cell_fill(c, "F5F7F8")

    doc.add_heading("31. Go-Live and Operating Checklists", level=1)
    checklists = {
        "Administrator go-live": [
            "Run migrations and installation health checks; confirm the active workflow, reports, workspace, cards, and charts.",
            "Assign roles and department User Permissions; test with representative users at each clearance level.",
            "Configure default currency, thresholds, email sender, retry limit, administrator alerts, and digest settings.",
            "Configure outgoing email and test both email and in-app reminders without exposing restricted records.",
            "Load master data, create sample leases, upload private documents, and validate reports before production import.",
        ],
        "Weekly operations": [
            "Review My Actions, Upcoming Expiries, Renewal Pipeline, Upcoming Payments, and document-expiry items.",
            "Investigate failed reminders and confirm retries or configuration corrections.",
            "Confirm newly paid schedules carry Payment Reference and Paid On; document every waiver reason.",
            "Confirm approved lifecycle requests have the correct current private document before execution.",
        ],
        "Monthly control review": [
            "Review role assignments, unrestricted users, department permissions, and Restricted lease access.",
            "Review Reminder Delivery rate and recurring failure causes.",
            "Reconcile Rent Exposure by currency and inspect approaching notice deadlines.",
            "Audit revision chains, superseded documents, and immutable reminder/decision history.",
        ],
    }
    for heading, items in checklists.items():
        doc.add_heading(heading, level=2)
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, "[ ] " + item)

    add_callout(doc, "End of knowledge base", "Keep this manual with the application release it documents. Rebuild or revise it whenever roles, DocTypes, workflow transitions, reports, scheduler jobs, or validation rules change.", PALE_TEAL)


def main():
    source = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    markdown_to_doc(doc, source)
    markdown_to_doc(doc, SUPPLEMENT)
    doctypes = user_doctypes()
    add_permissions_matrix(doc, doctypes)
    add_field_catalog(doc, doctypes)
    add_final_sections(doc)

    core = doc.core_properties
    core.title = "Rent Renewal Tracker Knowledge Base and User Manual"
    core.subject = "Complete application features, operation, roles, permissions, workflows and field reference"
    core.author = "Rent Renewal Tracker contributors"
    core.keywords = "lease, rent, renewal, Frappe, user manual, knowledge base"
    core.comments = "Generated from the application source as of 16 July 2026."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
